import io

from docx import Document

from app import clean_text, extract_resume_information, extract_text_from_docx, extract_text_from_txt


def test_txt_parser_and_cleaning():
    assert extract_text_from_txt(io.BytesIO(b"Maya\r\n\r\n\r\nPython")) == "Maya\n\nPython"


def test_docx_parser_includes_table():
    document = Document()
    document.add_paragraph("Nisha Rao")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Python | FastAPI"
    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    parsed = extract_text_from_docx(output)
    assert "Nisha Rao" in parsed
    assert "FastAPI" in parsed


def test_empty_text_handling_and_info():
    assert clean_text("  \n\n ") == ""
    info = extract_resume_information("Arjun Mehta\nExperience: 5 years\nEducation: B.Tech, 2020")
    assert info["candidate_name"] == "Arjun Mehta"
    assert info["experience_years"] == "5"
    assert info["graduation_year"] == "2020"
