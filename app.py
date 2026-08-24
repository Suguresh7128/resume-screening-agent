"""TalentRank AI: explainable resume screening with TF-IDF and Groq."""

from __future__ import annotations

import json
import os
import re
from typing import Any, BinaryIO

import pandas as pd
import streamlit as st
from docx import Document
from groq import Groq
from pypdf import PdfReader
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

DEFAULT_MODEL = "llama-3.3-70b-versatile"
MODEL_PREFERENCES = ["openai/gpt-oss-120b", "llama-4-scout-17b-16e-instruct", "qwen/qwen3-32b", "llama-3.3-70b-versatile"]
DEMO_DIR = os.path.join(os.path.dirname(__file__), "sample_data")
EXPORT_COLUMNS = ["Rank", "Candidate", "Final Score", "NLP Score", "LLM Score", "Experience", "Education", "Qualification", "Graduation Year", "Matched Skills", "Missing Skills", "Rationale", "Status"]


def clean_text(text: str) -> str:
    """Normalize whitespace without removing useful resume structure."""
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def extract_text_from_pdf(file: BinaryIO) -> str:
    """Extract text page by page from a PDF, tolerating textless pages."""
    pages = []
    for page in PdfReader(file).pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return clean_text("\n".join(pages))


def extract_text_from_docx(file: BinaryIO) -> str:
    """Extract paragraphs and table cells from a DOCX file."""
    document = Document(file)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return clean_text("\n".join(parts))


def extract_text_from_txt(file: BinaryIO) -> str:
    """Read UTF-8 text while tolerating imperfect uploaded files."""
    raw = file.read()
    if isinstance(raw, str):
        return clean_text(raw)
    return clean_text(raw.decode("utf-8", errors="replace"))


def extract_text_from_file(uploaded_file: Any) -> str:
    """Dispatch an uploaded file to its format-specific parser."""
    name = uploaded_file.name.lower()
    uploaded_file.seek(0)
    if name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    if name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    if name.endswith(".txt"):
        return extract_text_from_txt(uploaded_file)
    raise ValueError("Unsupported file type. Use PDF, DOCX, or TXT.")


def extract_resume_information(text: str) -> dict[str, Any]:
    """Perform conservative deterministic extraction used before LLM review."""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    name = "Not Specified"
    for line in lines[:8]:
        if 1 < len(line.split()) <= 5 and not re.search(r"@|resume|curriculum|developer|engineer|skills", line, re.I):
            name = line
            break
    years = re.search(r"(\d+(?:\.\d+)?)\+?\s+years?", text, re.I)
    grad = re.search(r"\b(19\d{2}|20\d{2})\b", text)
    education = next((line for line in lines if re.search(r"b\.?tech|bachelor|master|mca|m\.tech|degree|university|college", line, re.I)), "Not Specified")
    skill_line = next((line for line in lines if re.search(r"skills?\s*:", line, re.I)), "")
    skills = [s.strip() for s in re.split(r"[,|;]", skill_line.split(":", 1)[-1]) if s.strip()]
    return {"candidate_name": name, "experience_years": years.group(1) if years else "Not Specified", "education": education, "qualification": education, "graduation_year": grad.group(1) if grad else "Not Specified", "skills": skills}


def compute_tfidf_similarity(jd_text: str, resumes: list[dict[str, Any]]) -> list[float]:
    """Return JD-to-resume cosine similarities scaled to 0-100."""
    if not jd_text.strip() or not resumes:
        return [0.0 for _ in resumes]
    documents = [clean_text(jd_text)] + [clean_text(r.get("text", "")) for r in resumes]
    try:
        matrix = TfidfVectorizer(stop_words="english", ngram_range=(1, 2), max_features=5000).fit_transform(documents)
        return [round(float(value) * 100, 2) for value in cosine_similarity(matrix[0:1], matrix[1:]).flatten()]
    except ValueError:
        return [0.0 for _ in resumes]


def calculate_final_score(nlp_score: float, llm_score: float | None, nlp_weight: float) -> tuple[float, str]:
    """Combine successful scores, or clearly label an NLP-only fallback."""
    if llm_score is None:
        return round(nlp_score, 2), "NLP-only fallback (LLM failed)"
    return round(nlp_score * nlp_weight + llm_score * (1 - nlp_weight), 2), "Evaluated"


def status_for_score(score: float) -> str:
    if score >= 80:
        return "Strong Match"
    if score >= 65:
        return "Good Match"
    if score >= 45:
        return "Moderate Match"
    return "Weak Match"


def evaluate_candidate_with_llm(client: Groq, jd_text: str, resume_text: str, model: str) -> dict[str, Any]:
    """Ask Groq for factual structured evaluation; raise on API/JSON failure."""
    schema = {"candidate_name": "", "matched_skills": [], "missing_skills": [], "experience_years": "", "experience_details": "", "education": "", "qualification": "", "graduation_year": "", "skills": [], "fit_score": 0, "rationale": ""}
    system = """You are a careful technical recruiter. Read the complete resume and compare it to the JD. Do not invent information or confuse JD details with resume details. Missing information must be exactly 'Not Specified'. Education headings may be Education, Academic Background, Qualifications, Academic Qualifications, Degree, or Educational Qualification. Experience headings may be Experience, Work Experience, Professional Experience, Internship, or Employment; count relevant internships. Extract degree, institution, year, and GPA/CGPA when present. Return only valid JSON matching this schema: """ + json.dumps(schema)
    response = client.chat.completions.create(model=model, messages=[{"role": "system", "content": system}, {"role": "user", "content": f"JOB DESCRIPTION:\n{jd_text}\n\nCOMPLETE RESUME:\n{resume_text}\n\nReturn the JSON evaluation."}], temperature=0, response_format={"type": "json_object"})
    result = json.loads(response.choices[0].message.content)
    score = result.get("fit_score")
    if not isinstance(score, (int, float)) or not 0 <= float(score) <= 100:
        raise ValueError("The model returned an invalid fit_score.")
    result["fit_score"] = round(float(score), 2)
    return result


def discover_groq_models(api_key: str) -> list[str]:
    """Return chat-capable model IDs available to the current Groq key."""
    models = Groq(api_key=api_key).models.list()
    model_ids = []
    for model in models.data:
        model_id = getattr(model, "id", "")
        if model_id and "whisper" not in model_id and "safety" not in model_id and "guard" not in model_id:
            model_ids.append(model_id)
    return sorted(set(model_ids))


def resolve_groq_model(api_key: str, requested_model: str) -> tuple[str, list[str]]:
    """Use the requested model when available, otherwise select a preferred chat model."""
    available = discover_groq_models(api_key)
    if requested_model in available:
        return requested_model, available
    for preferred in MODEL_PREFERENCES:
        if preferred in available:
            return preferred, available
    if available:
        return available[0], available
    raise RuntimeError("Groq returned no available chat models for this API key.")


def rank_candidates(results: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Sort results strongest-first and assign stable ranks."""
    ordered = sorted(results, key=lambda item: item.get("Final Score", 0), reverse=True)
    for index, result in enumerate(ordered, 1):
        result["Rank"] = index
    return ordered


def create_csv(results: list[dict[str, Any]]) -> bytes:
    return pd.DataFrame(results, columns=EXPORT_COLUMNS).to_csv(index=False).encode("utf-8")


def create_json(results: list[dict[str, Any]]) -> bytes:
    return json.dumps([{column: result.get(column, "") for column in EXPORT_COLUMNS} for result in results], indent=2).encode("utf-8")


def load_demo_dataset() -> tuple[str, list[dict[str, Any]]]:
    with open(os.path.join(DEMO_DIR, "job_description.txt"), encoding="utf-8") as handle:
        jd = handle.read()
    resumes = []
    for filename in sorted(name for name in os.listdir(DEMO_DIR) if name.startswith("resume_") and name.endswith(".txt")):
        with open(os.path.join(DEMO_DIR, filename), encoding="utf-8") as handle:
            resumes.append({"filename": filename, "text": clean_text(handle.read()), "file_type": "TXT"})
    return jd, resumes


def build_result(resume: dict[str, Any], nlp_score: float, evaluation: dict[str, Any] | None, nlp_weight: float, error: str = "") -> dict[str, Any]:
    info = evaluation or extract_resume_information(resume["text"])
    llm_score = evaluation.get("fit_score") if evaluation else None
    final, evaluation_status = calculate_final_score(nlp_score, llm_score, nlp_weight)
    return {"Candidate": info.get("candidate_name") or resume["filename"], "Final Score": final, "NLP Score": nlp_score, "LLM Score": llm_score if llm_score is not None else "Unavailable", "Experience": info.get("experience_years", "Not Specified"), "Education": info.get("education", "Not Specified"), "Qualification": info.get("qualification", "Not Specified"), "Graduation Year": info.get("graduation_year", "Not Specified"), "Matched Skills": ", ".join(info.get("matched_skills", [])), "Missing Skills": ", ".join(info.get("missing_skills", [])), "Rationale": info.get("rationale", error or "LLM evaluation failed; score is NLP-only."), "Status": evaluation_status if evaluation_status != "Evaluated" else status_for_score(final), "LLM Status": "Successful" if evaluation else "Failed", "LLM Error": error, "filename": resume["filename"], "extracted_text": resume["text"], "Skills": info.get("skills", [])}


def run_screening(jd_text: str, resumes: list[dict[str, Any]], api_key: str | None, model: str, nlp_weight: float, progress=None) -> list[dict[str, Any]]:
    scores = compute_tfidf_similarity(jd_text, resumes)
    client = Groq(api_key=api_key) if api_key else None
    model_used = model
    model_notice = ""
    if client:
        try:
            model_used, available_models = resolve_groq_model(api_key, model)
            if model_used != model:
                model_notice = f"Requested model '{model}' was unavailable; automatically using '{model_used}'."
        except Exception as exc:
            model_notice = f"Model discovery failed: {str(exc)[:180]}"
    results = []
    for index, resume in enumerate(resumes):
        if progress:
            progress_value = (index + 1) / len(resumes)
            progress_text = f"Evaluating {resume['filename']} ({index + 1}/{len(resumes)})"
            if hasattr(progress, "progress"):
                progress.progress(progress_value, text=progress_text)
            else:
                progress(progress_value, progress_text)
        evaluation, error = None, ""
        if client:
            try:
                evaluation = evaluate_candidate_with_llm(client, jd_text, resume["text"], model_used)
            except Exception as exc:
                error = f"LLM evaluation failed: {str(exc)[:240]}"
        result = build_result(resume, scores[index], evaluation, nlp_weight, error)
        result["Model Used"] = model_used
        result["Model Notice"] = model_notice
        results.append(result)
    return rank_candidates(results)


def render_results(results: list[dict[str, Any]], threshold: float, nlp_weight: float) -> None:
    st.subheader("Ranked shortlist")
    model_notice = next((result.get("Model Notice") for result in results if result.get("Model Notice")), "")
    model_used = next((result.get("Model Used") for result in results if result.get("Model Used")), "")
    if model_notice:
        st.info(model_notice)
    if model_used and any(result.get("LLM Status") == "Successful" for result in results):
        st.success(f"LLM evaluation completed with {model_used}.")
    unavailable_model = any("model" in result.get("LLM Error", "").lower() and "404" in result.get("LLM Error", "") for result in results)
    if unavailable_model:
        st.error("The selected Groq model is unavailable for this API key. Use 'Discover available models' in the sidebar, select an available chat model, and run screening again.")
    average = sum(r["Final Score"] for r in results) / len(results)
    eligible = sum(r["Final Score"] >= threshold for r in results)
    best = results[0]["Candidate"] if results else "None"
    a, b, c, d = st.columns(4)
    a.metric("Candidates screened", len(results)); b.metric("Average score", f"{average:.1f}%"); c.metric("Top candidate", best); d.metric(f"Above {threshold:.0f}%", eligible)
    table = [{key: result.get(key, "") for key in EXPORT_COLUMNS} for result in results]
    st.dataframe(pd.DataFrame(table), use_container_width=True, hide_index=True)
    st.download_button("Download CSV", create_csv(results), "candidate_ranking_results.csv", "text/csv")
    st.download_button("Download JSON", create_json(results), "candidate_ranking_results.json", "application/json")
    chart = pd.DataFrame({"Candidate": [r["Candidate"] for r in results], "Final Score": [r["Final Score"] for r in results]}).set_index("Candidate")
    st.bar_chart(chart)
    comparison = pd.DataFrame({"NLP": [r["NLP Score"] for r in results], "LLM": [r["LLM Score"] if isinstance(r["LLM Score"], (int, float)) else 0 for r in results]}, index=[r["Candidate"] for r in results])
    st.line_chart(comparison)
    st.caption(f"Active formula: ({nlp_weight:.0%} x TF-IDF) + ({1 - nlp_weight:.0%} x LLM fit). Failed LLM evaluations remain NLP-only and are labeled.")
    for result in results:
        with st.expander(f"#{result['Rank']}  {result['Candidate']}  |  {result['Final Score']:.1f}%"):
            left, right = st.columns(2)
            left.write(f"**Experience:** {result['Experience']}"); left.write(f"**Education:** {result['Education']}"); left.write(f"**Qualification:** {result['Qualification']}"); left.write(f"**Graduation year:** {result['Graduation Year']}")
            right.write(f"**NLP:** {result['NLP Score']:.1f}%"); right.write(f"**LLM:** {result['LLM Score']}"); right.write(f"**Status:** {result['Status']} ({result['LLM Status']})")
            st.write(f"**Skills:** {', '.join(result['Skills']) or 'Not Specified'}")
            st.write(f"**Matched skills:** {result['Matched Skills'] or 'None identified'}")
            st.write(f"**Missing skills:** {result['Missing Skills'] or 'None identified'}")
            st.write(f"**AI reasoning:** {result['Rationale']}")
            if st.checkbox("Show extracted text", key=f"debug_{result['Rank']}"):
                st.text(result["extracted_text"])


st.set_page_config(page_title="TalentRank AI", page_icon="TR", layout="wide")
st.markdown("""<style>
.block-container{max-width:1400px;padding-top:2rem}
.hero{padding:1.7rem 0;border-bottom:1px solid #d9e2ec;margin-bottom:1.5rem;animation:rise .45s ease-out}
.hero h1{margin:.25rem 0 .15rem;color:#102a43;font-size:2.6rem;letter-spacing:0}
.hero p{color:#486581;margin:0;font-size:1.05rem}
.tag{color:#2f855a;font-weight:700;letter-spacing:.08em;text-transform:uppercase;font-size:.75rem}
[data-testid="stFileUploaderDropzone"]{border:1.5px dashed #7b9bb5;border-radius:12px;background:#f5f9fc;transition:all .2s ease}
[data-testid="stFileUploaderDropzone"]:hover{border-color:#2f855a;background:#eef8f1;transform:translateY(-1px);box-shadow:0 8px 22px rgba(16,42,67,.08)}
[data-testid="stMetric"]{background:rgba(255,255,255,.055);border:1px solid rgba(180,196,214,.22);border-radius:10px;padding:.8rem;animation:rise .35s ease-out;box-shadow:0 4px 14px rgba(0,0,0,.08)}
[data-testid="stMetricLabel"]{color:#9fb3c8}
[data-testid="stMetricValue"]{color:#f0f4f8}
div.stButton>button[kind="primary"]{background:#13795b;border:0;border-radius:8px;box-shadow:0 5px 14px rgba(19,121,91,.2);transition:transform .2s ease,box-shadow .2s ease}
div.stButton>button[kind="primary"]:hover{transform:translateY(-1px);box-shadow:0 8px 18px rgba(19,121,91,.28)}
@keyframes rise{from{opacity:0;transform:translateY(8px)}to{opacity:1;transform:translateY(0)}}
</style>""", unsafe_allow_html=True)
st.markdown('<div class="hero"><div class="tag">Recruiting intelligence</div><h1>TalentRank AI</h1><p>AI-Powered Resume Screening Agent</p></div>', unsafe_allow_html=True)

with st.sidebar:
    st.header("Configuration")
    environment_key = os.getenv("GROQ_API_KEY", "")
    entered_key = st.text_input("Groq API Key", type="password", help="Used only for this session.")
    api_key = environment_key or entered_key
    st.caption("Environment variable takes priority. Keys are never written to disk.")
    model = st.text_input("Groq model", value=DEFAULT_MODEL)
    if api_key and st.button("Discover available models"):
        try:
            st.session_state.available_models = discover_groq_models(api_key)
            if st.session_state.available_models:
                st.success(f"Found {len(st.session_state.available_models)} available models.")
            else:
                st.warning("Groq returned no chat models for this key.")
        except Exception as exc:
            st.error(f"Could not list Groq models: {str(exc)[:180]}")
    available_models = st.session_state.get("available_models", [])
    if available_models:
        model = st.selectbox("Available model", available_models, index=available_models.index(model) if model in available_models else 0)
    st.caption("Use model discovery if Groq reports that the selected model is unavailable.")
    nlp_weight = st.slider("NLP weight", 0.0, 1.0, 0.40, 0.05)
    st.metric("LLM weight", f"{1 - nlp_weight:.0%}")
    threshold = st.slider("Shortlist threshold", 0.0, 100.0, 60.0, 5.0)

if "jd" not in st.session_state:
    st.session_state.jd, st.session_state.demo_resumes = "", []
tab_screen, tab_rank, tab_insights, tab_method, tab_about = st.tabs(["Screen Candidates", "Rankings", "Candidate Insights", "Methodology", "About / Evaluation"])
with tab_screen:
    st.subheader("Build a shortlist")
    if st.button("Load Demo Dataset"):
        st.session_state.jd, st.session_state.demo_resumes = load_demo_dataset()
        st.success(f"Loaded {len(st.session_state.demo_resumes)} fictional sample resumes.")
    jd_text = st.text_area("Job description", value=st.session_state.jd, height=260, placeholder="Paste one job description here")
    uploads = st.file_uploader("Resumes (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], accept_multiple_files=True)
    resumes = st.session_state.demo_resumes if not uploads else []
    extraction_issues = []
    if uploads:
        for upload in uploads:
            try:
                text = extract_text_from_file(upload)
                if text:
                    resumes.append({"filename": upload.name, "text": text, "file_type": upload.name.rsplit(".", 1)[-1].upper()})
                else:
                    extraction_issues.append(f"{upload.name}: no text extracted")
            except Exception as exc:
                extraction_issues.append(f"{upload.name}: {str(exc)[:160]}")
    st.caption(f"{len(resumes)} resume(s) ready")
    if resumes:
        with st.expander("Extraction status"):
            for resume in resumes:
                st.write(f"Parsed successfully | {resume['filename']} | {resume['file_type']} | {len(resume['text']):,} characters")
            for issue in extraction_issues:
                st.warning(f"Extraction issue | {issue}")
    if st.button("Run AI Screening", type="primary", use_container_width=True):
        if not jd_text.strip(): st.error("Add a job description before screening.")
        elif not resumes: st.error("Upload resumes or load the demo dataset first.")
        elif not api_key: st.error("Add a Groq API key in the sidebar or set GROQ_API_KEY before running.")
        else:
            progress = st.progress(0, text="Parsing and evaluating resumes...")
            st.session_state.results = run_screening(jd_text, resumes, api_key, model, nlp_weight, progress)
            progress.progress(1.0, text="Screening complete")
            st.success("Screening complete. Open Rankings or Candidate Insights to review the shortlist.")
with tab_rank:
    results = st.session_state.get("results", [])
    if results: render_results(results, threshold, nlp_weight)
    else: st.info("Run a screening to populate rankings.")
with tab_insights:
    results = st.session_state.get("results", [])
    if results:
        for result in results:
            st.markdown(f"### {result['Candidate']}  ·  {result['Final Score']:.1f}%")
            st.write(result["Rationale"])
            st.progress(min(result["Final Score"] / 100, 1.0))
    else: st.info("Candidate insights appear after screening.")
with tab_method:
    st.subheader("Transparent hybrid scoring")
    st.markdown(f"**Final Score = ({nlp_weight:.0%} x TF-IDF similarity) + ({1 - nlp_weight:.0%} x LLM fit score)**")
    st.write("TF-IDF and cosine similarity measure explainable lexical overlap. Groq evaluates contextual skill, experience, education, gaps, and rationale. The recruiter controls the balance in the sidebar.")
    st.write("For an LLM failure, no artificial score is inserted: the candidate receives their NLP score as an explicitly labeled NLP-only fallback.")
with tab_about:
    st.subheader("Responsible screening")
    st.warning("This system is an assistive screening tool, not an autonomous hiring decision-maker. Evaluate job-relevant information only; do not score age, gender, religion, race, nationality, marital status, photographs, or other protected characteristics. A human recruiter makes the final decision.")
    st.write("Built for the ROOMAN AI Challenge, Category 1: HR & Recruitment. See README.md, docs/scoring_method.md, and docs/tradeoffs.md for reproducibility, decisions, and limitations.")
